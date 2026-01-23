"""
Document Generator Service for Polished Resume System.
Generates formatted PDF and DOCX resume documents.
"""

import io
import re
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass
from enum import Enum
from datetime import datetime

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, black
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    # Create placeholder values for when reportlab is not installed
    ParagraphStyle = None
    inch = 72  # 72 points per inch


class DocumentFormat(str, Enum):
    """Supported document formats."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"


class ResumeTemplate(str, Enum):
    """Resume template styles."""
    MODERN = "modern"
    CLASSIC = "classic"
    ATS_FRIENDLY = "ats_friendly"
    EXECUTIVE = "executive"
    MINIMAL = "minimal"


@dataclass
class ResumeData:
    """Structured resume data for document generation."""
    name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    summary: Optional[str] = None
    experience: Optional[List[Dict[str, Any]]] = None
    education: Optional[List[Dict[str, Any]]] = None
    skills: Optional[List[str]] = None
    certifications: Optional[List[str]] = None
    additional_sections: Optional[Dict[str, str]] = None


# Template color schemes
TEMPLATE_COLORS = {
    ResumeTemplate.MODERN: {
        "primary": "#2563eb",  # Blue
        "text": "#1f2937",
        "secondary": "#6b7280",
        "accent": "#3b82f6",
    },
    ResumeTemplate.CLASSIC: {
        "primary": "#1f2937",  # Dark gray
        "text": "#1f2937",
        "secondary": "#4b5563",
        "accent": "#1f2937",
    },
    ResumeTemplate.ATS_FRIENDLY: {
        "primary": "#000000",  # Black
        "text": "#000000",
        "secondary": "#374151",
        "accent": "#000000",
    },
    ResumeTemplate.EXECUTIVE: {
        "primary": "#1e3a5f",  # Navy
        "text": "#1f2937",
        "secondary": "#4b5563",
        "accent": "#2563eb",
    },
    ResumeTemplate.MINIMAL: {
        "primary": "#374151",  # Gray
        "text": "#1f2937",
        "secondary": "#6b7280",
        "accent": "#374151",
    },
}


class DocumentGenerator:
    """
    Generates formatted resume documents in various formats.

    Supports PDF, DOCX, TXT, and HTML output with multiple template styles.
    """

    def __init__(self):
        self.page_width = 8.5 * inch
        self.page_height = 11 * inch
        self.margin = 0.75 * inch

    def generate(
        self,
        resume_data: Union[ResumeData, Dict[str, Any]],
        format: DocumentFormat = DocumentFormat.PDF,
        template: ResumeTemplate = ResumeTemplate.MODERN,
    ) -> bytes:
        """
        Generate a formatted resume document.

        Args:
            resume_data: Resume content (ResumeData or dict)
            format: Output format (pdf, docx, txt, html)
            template: Template style

        Returns:
            Document bytes
        """
        # Convert dict to ResumeData if needed
        if isinstance(resume_data, dict):
            resume_data = ResumeData(**{
                k: v for k, v in resume_data.items()
                if k in ResumeData.__dataclass_fields__
            })

        if format == DocumentFormat.PDF:
            return self._generate_pdf(resume_data, template)
        elif format == DocumentFormat.DOCX:
            return self._generate_docx(resume_data, template)
        elif format == DocumentFormat.TXT:
            return self._generate_txt(resume_data)
        elif format == DocumentFormat.HTML:
            return self._generate_html(resume_data, template)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_pdf(self, data: ResumeData, template: ResumeTemplate) -> bytes:
        """Generate PDF document."""
        if not HAS_REPORTLAB:
            raise RuntimeError("reportlab not installed. Run: pip install reportlab")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=self.margin,
            leftMargin=self.margin,
            topMargin=self.margin,
            bottomMargin=self.margin,
        )

        colors = TEMPLATE_COLORS.get(template, TEMPLATE_COLORS[ResumeTemplate.MODERN])
        styles = self._get_pdf_styles(colors, template)
        story = []

        # Header / Contact
        story.extend(self._pdf_header(data, styles, colors))
        story.append(Spacer(1, 12))

        # Summary
        if data.summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", styles["section_header"]))
            story.append(Spacer(1, 6))
            story.append(Paragraph(data.summary, styles["body"]))
            story.append(Spacer(1, 12))

        # Experience
        if data.experience:
            story.append(Paragraph("EXPERIENCE", styles["section_header"]))
            story.append(Spacer(1, 6))
            for exp in data.experience:
                story.extend(self._pdf_experience_entry(exp, styles))
            story.append(Spacer(1, 8))

        # Education
        if data.education:
            story.append(Paragraph("EDUCATION", styles["section_header"]))
            story.append(Spacer(1, 6))
            for edu in data.education:
                story.extend(self._pdf_education_entry(edu, styles))
            story.append(Spacer(1, 8))

        # Skills
        if data.skills:
            story.append(Paragraph("SKILLS", styles["section_header"]))
            story.append(Spacer(1, 6))
            skills_text = " • ".join(data.skills)
            story.append(Paragraph(skills_text, styles["body"]))
            story.append(Spacer(1, 8))

        # Certifications
        if data.certifications:
            story.append(Paragraph("CERTIFICATIONS", styles["section_header"]))
            story.append(Spacer(1, 6))
            for cert in data.certifications:
                story.append(Paragraph(f"• {cert}", styles["body"]))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _get_pdf_styles(self, colors: Dict[str, str], template: ResumeTemplate) -> Dict[str, Any]:
        """Get PDF paragraph styles based on template."""
        styles = getSampleStyleSheet()

        return {
            "name": ParagraphStyle(
                "Name",
                parent=styles["Title"],
                fontSize=22 if template != ResumeTemplate.MINIMAL else 18,
                textColor=HexColor(colors["primary"]),
                spaceAfter=4,
                alignment=TA_CENTER if template in [ResumeTemplate.CLASSIC, ResumeTemplate.EXECUTIVE] else TA_LEFT,
            ),
            "contact": ParagraphStyle(
                "Contact",
                parent=styles["Normal"],
                fontSize=10,
                textColor=HexColor(colors["secondary"]),
                alignment=TA_CENTER if template in [ResumeTemplate.CLASSIC, ResumeTemplate.EXECUTIVE] else TA_LEFT,
            ),
            "section_header": ParagraphStyle(
                "SectionHeader",
                parent=styles["Heading2"],
                fontSize=12,
                textColor=HexColor(colors["primary"]),
                spaceBefore=8,
                spaceAfter=4,
                borderWidth=1 if template == ResumeTemplate.MODERN else 0,
                borderColor=HexColor(colors["accent"]),
                borderPadding=2,
            ),
            "company": ParagraphStyle(
                "Company",
                parent=styles["Normal"],
                fontSize=11,
                textColor=HexColor(colors["text"]),
                fontName="Helvetica-Bold",
            ),
            "role": ParagraphStyle(
                "Role",
                parent=styles["Normal"],
                fontSize=10,
                textColor=HexColor(colors["text"]),
                fontName="Helvetica-Oblique",
            ),
            "date": ParagraphStyle(
                "Date",
                parent=styles["Normal"],
                fontSize=10,
                textColor=HexColor(colors["secondary"]),
                alignment=TA_LEFT,
            ),
            "body": ParagraphStyle(
                "Body",
                parent=styles["Normal"],
                fontSize=10,
                textColor=HexColor(colors["text"]),
                alignment=TA_JUSTIFY if template == ResumeTemplate.EXECUTIVE else TA_LEFT,
                leading=14,
            ),
            "bullet": ParagraphStyle(
                "Bullet",
                parent=styles["Normal"],
                fontSize=10,
                textColor=HexColor(colors["text"]),
                leftIndent=12,
                leading=13,
            ),
        }

    def _pdf_header(self, data: ResumeData, styles: Dict, colors: Dict) -> List:
        """Generate PDF header section."""
        elements = []

        # Name
        elements.append(Paragraph(data.name.upper() if data.name else "CANDIDATE", styles["name"]))

        # Contact info
        contact_parts = []
        if data.email:
            contact_parts.append(data.email)
        if data.phone:
            contact_parts.append(data.phone)
        if data.location:
            contact_parts.append(data.location)
        if data.linkedin:
            contact_parts.append(data.linkedin)

        if contact_parts:
            contact_text = " | ".join(contact_parts)
            elements.append(Paragraph(contact_text, styles["contact"]))

        return elements

    def _pdf_experience_entry(self, exp: Dict, styles: Dict) -> List:
        """Generate PDF experience entry."""
        elements = []

        company = exp.get("company", "Company")
        title = exp.get("title", "Role")
        dates = exp.get("dates", "")
        location = exp.get("location", "")
        bullets = exp.get("bullets", [])

        # Company and dates on same line
        company_text = f"<b>{company}</b>"
        if location:
            company_text += f" | {location}"

        elements.append(Paragraph(company_text, styles["company"]))

        role_text = f"<i>{title}</i>"
        if dates:
            role_text += f"  |  {dates}"
        elements.append(Paragraph(role_text, styles["role"]))
        elements.append(Spacer(1, 4))

        # Bullets
        for bullet in bullets:
            elements.append(Paragraph(f"• {bullet}", styles["bullet"]))

        elements.append(Spacer(1, 8))
        return elements

    def _pdf_education_entry(self, edu: Dict, styles: Dict) -> List:
        """Generate PDF education entry."""
        elements = []

        school = edu.get("school", "")
        degree = edu.get("degree", "")
        field = edu.get("field", "")
        year = edu.get("year", "")
        gpa = edu.get("gpa", "")

        school_text = f"<b>{school}</b>"
        if year:
            school_text += f"  |  {year}"
        elements.append(Paragraph(school_text, styles["company"]))

        degree_text = degree
        if field:
            degree_text += f" in {field}"
        if gpa:
            degree_text += f" (GPA: {gpa})"
        if degree_text:
            elements.append(Paragraph(degree_text, styles["body"]))

        elements.append(Spacer(1, 6))
        return elements

    def _generate_docx(self, data: ResumeData, template: ResumeTemplate) -> bytes:
        """Generate DOCX document."""
        if not HAS_DOCX:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        doc = Document()
        colors = TEMPLATE_COLORS.get(template, TEMPLATE_COLORS[ResumeTemplate.MODERN])

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Name
        name_para = doc.add_heading(data.name.upper() if data.name else "CANDIDATE", level=0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if template in [
            ResumeTemplate.CLASSIC, ResumeTemplate.EXECUTIVE
        ] else WD_ALIGN_PARAGRAPH.LEFT

        # Contact
        contact_parts = []
        if data.email:
            contact_parts.append(data.email)
        if data.phone:
            contact_parts.append(data.phone)
        if data.location:
            contact_parts.append(data.location)
        if data.linkedin:
            contact_parts.append(data.linkedin)

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if template in [
                ResumeTemplate.CLASSIC, ResumeTemplate.EXECUTIVE
            ] else WD_ALIGN_PARAGRAPH.LEFT

        # Summary
        if data.summary:
            doc.add_heading("PROFESSIONAL SUMMARY", level=1)
            doc.add_paragraph(data.summary)

        # Experience
        if data.experience:
            doc.add_heading("EXPERIENCE", level=1)
            for exp in data.experience:
                self._docx_experience_entry(doc, exp)

        # Education
        if data.education:
            doc.add_heading("EDUCATION", level=1)
            for edu in data.education:
                self._docx_education_entry(doc, edu)

        # Skills
        if data.skills:
            doc.add_heading("SKILLS", level=1)
            doc.add_paragraph(" • ".join(data.skills))

        # Certifications
        if data.certifications:
            doc.add_heading("CERTIFICATIONS", level=1)
            for cert in data.certifications:
                doc.add_paragraph(f"• {cert}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _docx_experience_entry(self, doc: Document, exp: Dict):
        """Add experience entry to DOCX."""
        company = exp.get("company", "Company")
        title = exp.get("title", "Role")
        dates = exp.get("dates", "")
        location = exp.get("location", "")
        bullets = exp.get("bullets", [])

        # Company line
        company_para = doc.add_paragraph()
        run = company_para.add_run(company)
        run.bold = True
        if location:
            company_para.add_run(f" | {location}")
        if dates:
            company_para.add_run(f"  |  {dates}")

        # Title
        title_para = doc.add_paragraph()
        run = title_para.add_run(title)
        run.italic = True

        # Bullets
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    def _docx_education_entry(self, doc: Document, edu: Dict):
        """Add education entry to DOCX."""
        school = edu.get("school", "")
        degree = edu.get("degree", "")
        field = edu.get("field", "")
        year = edu.get("year", "")
        gpa = edu.get("gpa", "")

        # School line
        school_para = doc.add_paragraph()
        run = school_para.add_run(school)
        run.bold = True
        if year:
            school_para.add_run(f"  |  {year}")

        # Degree
        if degree:
            degree_text = degree
            if field:
                degree_text += f" in {field}"
            if gpa:
                degree_text += f" (GPA: {gpa})"
            doc.add_paragraph(degree_text)

    def _generate_txt(self, data: ResumeData) -> bytes:
        """Generate plain text document."""
        lines = []

        # Header
        lines.append(data.name.upper() if data.name else "CANDIDATE")
        lines.append("=" * len(lines[0]))

        contact_parts = []
        if data.email:
            contact_parts.append(data.email)
        if data.phone:
            contact_parts.append(data.phone)
        if data.location:
            contact_parts.append(data.location)
        if data.linkedin:
            contact_parts.append(data.linkedin)
        if contact_parts:
            lines.append(" | ".join(contact_parts))

        lines.append("")

        # Summary
        if data.summary:
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 20)
            lines.append(data.summary)
            lines.append("")

        # Experience
        if data.experience:
            lines.append("EXPERIENCE")
            lines.append("-" * 20)
            for exp in data.experience:
                company = exp.get("company", "")
                title = exp.get("title", "")
                dates = exp.get("dates", "")
                location = exp.get("location", "")
                bullets = exp.get("bullets", [])

                header = company
                if location:
                    header += f" | {location}"
                if dates:
                    header += f" | {dates}"
                lines.append(header)
                lines.append(f"  {title}")
                for bullet in bullets:
                    lines.append(f"    • {bullet}")
                lines.append("")

        # Education
        if data.education:
            lines.append("EDUCATION")
            lines.append("-" * 20)
            for edu in data.education:
                school = edu.get("school", "")
                degree = edu.get("degree", "")
                field = edu.get("field", "")
                year = edu.get("year", "")

                header = school
                if year:
                    header += f" | {year}"
                lines.append(header)
                if degree:
                    degree_text = degree
                    if field:
                        degree_text += f" in {field}"
                    lines.append(f"  {degree_text}")
                lines.append("")

        # Skills
        if data.skills:
            lines.append("SKILLS")
            lines.append("-" * 20)
            lines.append(" • ".join(data.skills))
            lines.append("")

        # Certifications
        if data.certifications:
            lines.append("CERTIFICATIONS")
            lines.append("-" * 20)
            for cert in data.certifications:
                lines.append(f"• {cert}")

        return "\n".join(lines).encode("utf-8")

    def _generate_html(self, data: ResumeData, template: ResumeTemplate) -> bytes:
        """Generate HTML document."""
        colors = TEMPLATE_COLORS.get(template, TEMPLATE_COLORS[ResumeTemplate.MODERN])

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{data.name or 'Resume'}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            color: {colors['text']};
            line-height: 1.5;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 0.75in;
            background: white;
        }}
        .header {{
            text-align: {'center' if template in [ResumeTemplate.CLASSIC, ResumeTemplate.EXECUTIVE] else 'left'};
            margin-bottom: 1rem;
        }}
        h1 {{
            color: {colors['primary']};
            font-size: 1.75rem;
            margin-bottom: 0.5rem;
        }}
        .contact {{
            color: {colors['secondary']};
            font-size: 0.9rem;
        }}
        h2 {{
            color: {colors['primary']};
            font-size: 1rem;
            border-bottom: {'2px solid ' + colors['accent'] if template == ResumeTemplate.MODERN else 'none'};
            padding-bottom: 0.25rem;
            margin: 1rem 0 0.5rem;
            text-transform: uppercase;
        }}
        .experience-item, .education-item {{
            margin-bottom: 0.75rem;
        }}
        .company {{
            font-weight: bold;
        }}
        .role {{
            font-style: italic;
        }}
        .date {{
            color: {colors['secondary']};
        }}
        ul {{
            margin-left: 1.25rem;
            margin-top: 0.25rem;
        }}
        li {{
            margin-bottom: 0.25rem;
        }}
        .skills {{
            display: flex;
            flex-wrap: wrap;
            gap: 0.5rem;
        }}
        .skill {{
            background: #f3f4f6;
            padding: 0.25rem 0.5rem;
            border-radius: 0.25rem;
            font-size: 0.875rem;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{data.name.upper() if data.name else 'CANDIDATE'}</h1>
        <p class="contact">
            {' | '.join(filter(None, [data.email, data.phone, data.location, data.linkedin]))}
        </p>
    </div>
"""

        # Summary
        if data.summary:
            html += f"""
    <section>
        <h2>Professional Summary</h2>
        <p>{data.summary}</p>
    </section>
"""

        # Experience
        if data.experience:
            html += """
    <section>
        <h2>Experience</h2>
"""
            for exp in data.experience:
                company = exp.get("company", "")
                title = exp.get("title", "")
                dates = exp.get("dates", "")
                location = exp.get("location", "")
                bullets = exp.get("bullets", [])

                html += f"""
        <div class="experience-item">
            <p><span class="company">{company}</span>{f' | {location}' if location else ''}</p>
            <p><span class="role">{title}</span>{f' | <span class="date">{dates}</span>' if dates else ''}</p>
"""
                if bullets:
                    html += "            <ul>\n"
                    for bullet in bullets:
                        html += f"                <li>{bullet}</li>\n"
                    html += "            </ul>\n"
                html += "        </div>\n"
            html += "    </section>\n"

        # Education
        if data.education:
            html += """
    <section>
        <h2>Education</h2>
"""
            for edu in data.education:
                school = edu.get("school", "")
                degree = edu.get("degree", "")
                field = edu.get("field", "")
                year = edu.get("year", "")

                html += f"""
        <div class="education-item">
            <p><span class="company">{school}</span>{f' | <span class="date">{year}</span>' if year else ''}</p>
"""
                if degree:
                    html += f"            <p>{degree}{f' in {field}' if field else ''}</p>\n"
                html += "        </div>\n"
            html += "    </section>\n"

        # Skills
        if data.skills:
            html += """
    <section>
        <h2>Skills</h2>
        <div class="skills">
"""
            for skill in data.skills:
                html += f'            <span class="skill">{skill}</span>\n'
            html += """        </div>
    </section>
"""

        # Certifications
        if data.certifications:
            html += """
    <section>
        <h2>Certifications</h2>
        <ul>
"""
            for cert in data.certifications:
                html += f"            <li>{cert}</li>\n"
            html += """        </ul>
    </section>
"""

        html += """
</body>
</html>
"""

        return html.encode("utf-8")

    def parse_resume_text(self, resume_text: str, extracted_data: Optional[Dict] = None) -> ResumeData:
        """
        Parse resume text into structured ResumeData.

        Args:
            resume_text: Raw resume text
            extracted_data: Pre-extracted data from resume processing

        Returns:
            ResumeData object
        """
        data = extracted_data or {}

        # Parse experience from text if not in extracted_data
        experience = data.get("experience", [])
        if not experience and resume_text:
            experience = self._parse_experience_from_text(resume_text)

        # Parse education
        education = data.get("education", [])
        if not education and resume_text:
            education = self._parse_education_from_text(resume_text)

        return ResumeData(
            name=data.get("name", ""),
            email=data.get("email"),
            phone=data.get("phone"),
            location=data.get("location"),
            linkedin=data.get("linkedin"),
            summary=data.get("summary"),
            experience=experience,
            education=education,
            skills=data.get("skills", []),
            certifications=data.get("certifications", []),
        )

    def _parse_experience_from_text(self, text: str) -> List[Dict]:
        """Simple experience parsing from text."""
        experience = []

        # Look for experience section
        exp_match = re.search(
            r"(?:EXPERIENCE|WORK EXPERIENCE|PROFESSIONAL EXPERIENCE)(.*?)(?:EDUCATION|SKILLS|$)",
            text,
            re.IGNORECASE | re.DOTALL
        )

        if exp_match:
            exp_text = exp_match.group(1)
            # Split by common patterns
            entries = re.split(r"\n(?=[A-Z][A-Za-z\s]+(?:Inc|LLC|Corp|Company|Co\.)?\s*[\|])", exp_text)

            for entry in entries:
                if len(entry.strip()) < 20:
                    continue

                lines = entry.strip().split("\n")
                if not lines:
                    continue

                # First line is usually company
                company_line = lines[0].strip()
                company_parts = re.split(r"\s*[\|]\s*", company_line)
                company = company_parts[0] if company_parts else ""
                location = company_parts[1] if len(company_parts) > 1 else ""

                # Second line is usually title and dates
                title = ""
                dates = ""
                if len(lines) > 1:
                    title_line = lines[1].strip()
                    title_parts = re.split(r"\s*[\|]\s*", title_line)
                    title = title_parts[0] if title_parts else ""
                    if len(title_parts) > 1:
                        dates = title_parts[-1]

                # Remaining lines are bullets
                bullets = []
                for line in lines[2:]:
                    line = line.strip().lstrip("•-* ")
                    if len(line) > 10:
                        bullets.append(line)

                if company or title:
                    experience.append({
                        "company": company,
                        "title": title,
                        "location": location,
                        "dates": dates,
                        "bullets": bullets[:6],  # Limit to 6 bullets
                    })

        return experience[:5]  # Limit to 5 positions

    def _parse_education_from_text(self, text: str) -> List[Dict]:
        """Simple education parsing from text."""
        education = []

        edu_match = re.search(
            r"(?:EDUCATION|ACADEMIC)(.*?)(?:SKILLS|CERTIFICATIONS|$)",
            text,
            re.IGNORECASE | re.DOTALL
        )

        if edu_match:
            edu_text = edu_match.group(1)

            # Look for degree patterns
            degree_patterns = [
                r"(Bachelor|Master|PhD|MBA|BS|MS|BA|MA)[^,\n]*(?:,|\n|from)\s*([A-Za-z\s]+(?:University|College|Institute))",
                r"([A-Za-z\s]+(?:University|College|Institute))[^,\n]*(?:,|\n)\s*(Bachelor|Master|PhD|MBA|BS|MS|BA|MA)",
            ]

            for pattern in degree_patterns:
                matches = re.findall(pattern, edu_text, re.IGNORECASE)
                for match in matches:
                    degree = match[0] if "University" not in match[0] else match[1]
                    school = match[1] if "University" in match[1] else match[0]

                    # Try to find year
                    year_match = re.search(r"20\d{2}|19\d{2}", edu_text)
                    year = year_match.group(0) if year_match else ""

                    education.append({
                        "school": school.strip(),
                        "degree": degree.strip(),
                        "year": year,
                    })

        return education[:3]


# Singleton instance
_doc_generator: Optional[DocumentGenerator] = None


def get_generator() -> DocumentGenerator:
    """Get the document generator singleton."""
    global _doc_generator
    if _doc_generator is None:
        _doc_generator = DocumentGenerator()
    return _doc_generator


def generate_document(
    resume_data: Union[ResumeData, Dict[str, Any]],
    format: str = "pdf",
    template: str = "modern",
) -> bytes:
    """
    Convenience function to generate a resume document.

    Args:
        resume_data: Resume content
        format: Output format (pdf, docx, txt, html)
        template: Template style

    Returns:
        Document bytes
    """
    generator = get_generator()
    return generator.generate(
        resume_data=resume_data,
        format=DocumentFormat(format),
        template=ResumeTemplate(template),
    )
